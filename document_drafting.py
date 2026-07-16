"""
document_drafting.py

Herramienta de redaccion asistida de documentos legales, por terminal.

Flujo:
  1. El usuario elige un tipo de documento de un menu numerado.
  2. Se piden los datos necesarios para ese tipo de documento (con
     validaciones basicas de monto, correo y fecha).
  3. Se genera un borrador con la IA configurada en config.py, usando
     lenguaje juridico formal y la estructura estandar del documento.
  4. El usuario puede pedir cambios y se regenera el documento las veces
     que haga falta.
  5. El resultado final se exporta a .txt y .docx (con formato profesional)
     y se registra en historial.json.

IMPORTANTE: los documentos generados son borradores de apoyo. No sustituyen
la revision de un abogado antes de presentarlos ante una autoridad.

Uso:
    python document_drafting.py
"""

import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import config
from legal_research import ask_llm

SALIDA_DIR = Path(__file__).resolve().parent / "documentos_generados"
HISTORIAL_PATH = Path(__file__).resolve().parent / "historial.json"

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")

DRAFTING_SYSTEM_PROMPT = (
    "Eres un abogado experto en redaccion de documentos legales para Mexico. "
    "Redactas en espanol, con lenguaje juridico formal y preciso, siguiendo "
    "la estructura estandar del tipo de documento solicitado: encabezado, "
    "cuerpo (antecedentes/hechos, fundamentos de derecho, peticiones u "
    "objeto segun corresponda), lugar y fecha, y espacio de firma. Citas "
    "articulos de ley genericos y razonables cuando sea pertinente (ej. "
    "Codigo Civil, Codigo de Procedimientos Civiles), aclarando que son una "
    "referencia orientativa y no sustituyen la revision de un abogado. No "
    "inventes hechos ni datos que no se te hayan proporcionado; usa "
    "exactamente los datos entregados."
)


# ---------------------------------------------------------------------------
# Catalogo de tipos de documento y sus campos
# ---------------------------------------------------------------------------
# Cada campo es (clave, etiqueta, tipo, opcional). Tipos soportados:
# "text", "multiline", "money", "date", "email".

TIPOS_DOCUMENTO = {
    "1": {
        "nombre": "Carta de demanda (incumplimiento de contrato / daños)",
        "enfoque": (
            "Carta de demanda formal ante juzgado civil por incumplimiento "
            "de contrato y/o daños, con antecedentes, fundamentos de "
            "derecho y petitorio."
        ),
        "campos": [
            ("actor_nombre", "Nombre completo del actor (quien demanda)", "text", False),
            ("actor_domicilio", "Domicilio completo del actor", "text", False),
            ("demandado_nombre", "Nombre completo del demandado", "text", False),
            ("demandado_domicilio", "Domicilio completo del demandado", "text", False),
            ("fecha_contrato", "Fecha del contrato (DD/MM/YYYY)", "date", False),
            ("fecha_incumplimiento", "Fecha del incumplimiento (DD/MM/YYYY)", "date", False),
            ("monto_reclamado", "Monto reclamado (solo numero, ej. 50000)", "money", False),
            ("hechos", "Hechos relevantes (resumen de lo sucedido)", "multiline", False),
            ("juzgado", "Juzgado o autoridad competente", "text", False),
            ("expediente", "Numero de expediente", "text", True),
        ],
    },
    "2": {
        "nombre": "Contrato de servicios profesionales",
        "enfoque": (
            "Contrato de prestación de servicios profesionales entre un "
            "prestador y un cliente, con objeto, honorarios, vigencia y "
            "cláusulas estándar."
        ),
        "campos": [
            ("prestador_nombre", "Nombre completo del prestador de servicios", "text", False),
            ("prestador_domicilio", "Domicilio del prestador de servicios", "text", False),
            ("prestador_rfc", "RFC del prestador de servicios", "text", True),
            ("cliente_nombre", "Nombre completo del cliente", "text", False),
            ("cliente_domicilio", "Domicilio del cliente", "text", False),
            ("cliente_rfc", "RFC/CURP del cliente", "text", True),
            ("cliente_email", "Correo electrónico del cliente", "email", True),
            ("objeto_servicio", "Descripción del servicio a prestar", "multiline", False),
            ("monto_honorarios", "Monto de los honorarios (solo numero)", "money", False),
            ("fecha_inicio", "Fecha de inicio (DD/MM/YYYY)", "date", False),
            ("fecha_fin", "Fecha de término (DD/MM/YYYY)", "date", True),
            ("lugar_firma", "Lugar de firma del contrato", "text", False),
        ],
    },
    "3": {
        "nombre": "Escrito de contestación de demanda",
        "enfoque": (
            "Escrito de contestación de demanda dirigido al juzgado, "
            "respondiendo los hechos y oponiendo excepciones y defensas."
        ),
        "campos": [
            ("demandado_nombre", "Nombre completo del demandado (quien contesta)", "text", False),
            ("demandado_domicilio", "Domicilio del demandado", "text", False),
            ("actor_nombre", "Nombre completo del actor", "text", False),
            ("expediente", "Numero de expediente", "text", False),
            ("juzgado", "Juzgado o autoridad ante quien se contesta", "text", False),
            ("fecha_demanda_recibida", "Fecha en que se recibió la demanda (DD/MM/YYYY)", "date", False),
            ("hechos_contestacion", "Respuesta punto por punto a los hechos de la demanda", "multiline", False),
            ("excepciones", "Excepciones y defensas que se hacen valer", "multiline", False),
        ],
    },
    "4": {
        "nombre": "Convenio de pago",
        "enfoque": (
            "Convenio de pago en parcialidades entre acreedor y deudor, "
            "con monto, plazos y consecuencias por incumplimiento."
        ),
        "campos": [
            ("acreedor_nombre", "Nombre completo del acreedor", "text", False),
            ("deudor_nombre", "Nombre completo del deudor", "text", False),
            ("monto_adeudado", "Monto total adeudado (solo numero)", "money", False),
            ("numero_pagos", "Número de pagos/parcialidades acordadas", "text", False),
            ("fecha_primer_pago", "Fecha del primer pago (DD/MM/YYYY)", "date", False),
            ("hechos", "Origen de la deuda (resumen)", "multiline", False),
            ("lugar_firma", "Lugar de firma del convenio", "text", False),
        ],
    },
    "5": {
        "nombre": "Carta de requerimiento de pago",
        "enfoque": (
            "Carta de requerimiento formal de pago, notificando la deuda y "
            "el plazo otorgado antes de iniciar acciones legales."
        ),
        "campos": [
            ("acreedor_nombre", "Nombre completo de quien requiere el pago", "text", False),
            ("deudor_nombre", "Nombre completo del deudor", "text", False),
            ("deudor_domicilio", "Domicilio del deudor", "text", False),
            ("monto_adeudado", "Monto adeudado (solo numero)", "money", False),
            ("fecha_vencimiento", "Fecha de vencimiento de la deuda (DD/MM/YYYY)", "date", False),
            ("plazo_dias", "Plazo en dias que se otorga para pagar", "text", False),
            ("hechos", "Origen de la deuda (resumen)", "multiline", False),
        ],
    },
    "6": {
        "nombre": "Poder especial",
        "enfoque": (
            "Poder especial otorgado ante notario o autoridad competente, "
            "especificando facultades concretas y limitadas."
        ),
        "campos": [
            ("otorgante_nombre", "Nombre completo de quien otorga el poder", "text", False),
            ("otorgante_domicilio", "Domicilio del otorgante", "text", False),
            ("apoderado_nombre", "Nombre completo del apoderado", "text", False),
            ("facultades", "Facultades que se otorgan (qué puede hacer el apoderado)", "multiline", False),
            ("fecha_otorgamiento", "Fecha de otorgamiento (DD/MM/YYYY)", "date", False),
            ("lugar_otorgamiento", "Lugar de otorgamiento", "text", False),
        ],
    },
}


# ---------------------------------------------------------------------------
# Numero a letras (para montos en pesos)
# ---------------------------------------------------------------------------

_UNIDADES = ["", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
_DIECIS = [
    "diez", "once", "doce", "trece", "catorce", "quince",
    "dieciseis", "diecisiete", "dieciocho", "diecinueve",
]
_DECENAS = ["", "", "veinte", "treinta", "cuarenta", "cincuenta", "sesenta", "setenta", "ochenta", "noventa"]
_CENTENAS = [
    "", "ciento", "doscientos", "trescientos", "cuatrocientos", "quinientos",
    "seiscientos", "setecientos", "ochocientos", "novecientos",
]


def _tres_digitos_a_letras(numero: int) -> str:
    if numero == 0:
        return ""
    if numero == 100:
        return "cien"

    centena, resto = divmod(numero, 100)
    decena, unidad = divmod(resto, 10)

    partes = []
    if centena:
        partes.append(_CENTENAS[centena])

    if resto == 0:
        pass
    elif resto < 10:
        partes.append(_UNIDADES[unidad])
    elif resto < 20:
        partes.append(_DIECIS[resto - 10])
    elif resto == 20:
        partes.append("veinte")
    elif unidad == 0:
        partes.append(_DECENAS[decena])
    else:
        partes.append(f"{_DECENAS[decena]} y {_UNIDADES[unidad]}")

    return " ".join(partes)


def numero_a_letras(monto: float) -> str:
    """Convierte un monto a su representacion en letras para pesos
    mexicanos, ej. 1500.5 -> 'Mil quinientos pesos 50/100 M.N.'"""
    entero = int(monto)
    centavos = round((monto - entero) * 100)

    if entero == 0:
        texto_entero = "cero"
    else:
        millones, resto = divmod(entero, 1_000_000)
        miles, unidades_resto = divmod(resto, 1000)

        partes = []
        if millones:
            partes.append("un millon" if millones == 1 else f"{_tres_digitos_a_letras(millones)} millones")
        if miles:
            partes.append("mil" if miles == 1 else f"{_tres_digitos_a_letras(miles)} mil")
        if unidades_resto or not partes:
            partes.append(_tres_digitos_a_letras(unidades_resto))

        texto_entero = " ".join(p for p in partes if p)

    texto = f"{texto_entero} pesos {centavos:02d}/100 M.N."
    return texto[0].upper() + texto[1:]


# ---------------------------------------------------------------------------
# Recoleccion de datos con validacion
# ---------------------------------------------------------------------------

def validar_monto(valor_texto: str) -> float:
    """Valida y convierte un texto de monto a float. Lanza ValueError con un
    mensaje legible si no es valido. Reusada tanto por la CLI como por la UI."""
    try:
        monto = float(valor_texto.replace("$", "").replace(",", "").strip())
    except ValueError:
        raise ValueError("Monto invalido. Escribe solo numeros (ej. 50000 o 50000.00).")
    if monto <= 0:
        raise ValueError("El monto debe ser mayor a cero.")
    return monto


def validar_email(valor: str) -> str:
    """Valida un correo electronico. Lanza ValueError si no es valido."""
    if not EMAIL_RE.match(valor):
        raise ValueError("Correo invalido. Ejemplo: nombre@dominio.com")
    return valor


def pedir_dato(etiqueta: str, tipo: str = "text", opcional: bool = False):
    """Pide un dato por terminal y lo valida segun su tipo. Repite hasta
    obtener un valor valido (o vacio, si el campo es opcional)."""
    sufijo = " (opcional, Enter para omitir)" if opcional else ""
    while True:
        valor = input(f"{etiqueta}{sufijo}: ").strip()

        if not valor:
            if opcional:
                return ""
            print("  Este dato es obligatorio.")
            continue

        if tipo == "money":
            try:
                return validar_monto(valor)
            except ValueError as exc:
                print(f"  {exc}")
                continue

        if tipo == "email":
            try:
                return validar_email(valor)
            except ValueError as exc:
                print(f"  {exc}")
                continue

        if tipo == "date":
            try:
                datetime.strptime(valor, "%d/%m/%Y")
            except ValueError:
                print("  Fecha invalida. Usa el formato DD/MM/YYYY (ej. 15/03/2026).")
                continue
            return valor

        return valor


def pedir_texto_multilinea(etiqueta: str) -> str:
    """Pide un bloque de texto libre; el usuario termina dejando una linea
    vacia despues de haber escrito algo."""
    print(f"{etiqueta} (escribe tu texto; deja una linea vacía para terminar):")
    lineas: list[str] = []
    while True:
        linea = input()
        if not linea:
            if lineas:
                break
            continue
        lineas.append(linea)
    return "\n".join(lineas)


def elegir_tipo_documento() -> str | None:
    print("\nTipos de documento disponibles:\n")
    for clave, info in TIPOS_DOCUMENTO.items():
        print(f"  {clave}. {info['nombre']}")
    print("  0. Cancelar\n")

    while True:
        eleccion = input("Elige una opcion: ").strip()
        if eleccion == "0":
            return None
        if eleccion in TIPOS_DOCUMENTO:
            return eleccion
        print("Opcion invalida, intenta de nuevo.")


def recolectar_datos(tipo_info: dict) -> dict:
    datos = {}
    for clave, etiqueta, tipo_campo, opcional in tipo_info["campos"]:
        if tipo_campo == "multiline":
            datos[clave] = pedir_texto_multilinea(etiqueta)
        else:
            datos[clave] = pedir_dato(etiqueta, tipo_campo, opcional)
    return datos


def formatear_datos_para_prompt(tipo_info: dict, datos: dict) -> list[tuple[str, str]]:
    """Convierte los datos crudos en pares (etiqueta, texto) listos para
    incluir en el prompt, formateando montos con letra y numero."""
    resultado = []
    for clave, etiqueta, tipo_campo, _ in tipo_info["campos"]:
        valor = datos[clave]
        if valor == "":
            continue
        if tipo_campo == "money":
            valor_texto = f"${valor:,.2f} ({numero_a_letras(valor)})"
        else:
            valor_texto = str(valor)
        resultado.append((etiqueta, valor_texto))
    return resultado


def extraer_partes(tipo_info: dict, datos: dict) -> list[str]:
    """Extrae los nombres de las partes involucradas (heuristica: cualquier
    campo cuya clave contenga 'nombre'), para guardarlos en el historial."""
    return [
        datos[clave]
        for clave, _, _, _ in tipo_info["campos"]
        if "nombre" in clave and datos[clave]
    ]


# ---------------------------------------------------------------------------
# Generacion con IA
# ---------------------------------------------------------------------------

def construir_prompt(tipo_nombre: str, enfoque: str, datos_formateados: list[tuple[str, str]]) -> str:
    lineas_datos = "\n".join(f"- {etiqueta}: {valor}" for etiqueta, valor in datos_formateados)
    return (
        f"Redacta un(a) {tipo_nombre}.\n\n"
        f"Enfoque: {enfoque}\n\n"
        f"Datos proporcionados:\n{lineas_datos}\n\n"
        "Genera el documento completo y listo para usar, con encabezado, "
        "cuerpo, peticiones u objeto según corresponda, lugar y fecha, y "
        "espacio para firma."
    )


def construir_prompt_cambios(borrador_previo: str, instrucciones: str) -> str:
    return (
        f"Este es el borrador actual del documento:\n\n{borrador_previo}\n\n"
        f"Aplica los siguientes cambios solicitados:\n{instrucciones}\n\n"
        "Devuelve el documento completo revisado (no solo los cambios)."
    )


def generar_borrador(prompt: str, provider_config: dict) -> str | None:
    try:
        return ask_llm(prompt, provider_config, DRAFTING_SYSTEM_PROMPT)
    except Exception as exc:
        print(f"\nNo se pudo generar el documento: {exc}")
        print("Verifica tu conexion y la configuracion del proveedor de LLM en .env.\n")
        return None


# ---------------------------------------------------------------------------
# Exportacion: .txt y .docx
# ---------------------------------------------------------------------------

def _agregar_numero_pagina(seccion) -> None:
    """Inserta un campo de numero de pagina (PAGE) en el pie de pagina."""
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = parrafo.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def exportar_docx(destino, titulo: str, cuerpo: str) -> None:
    """destino puede ser una Path/str (uso por CLI) o un stream en memoria
    tipo io.BytesIO (uso desde la UI, para ofrecer descarga sin escribir a
    disco). python-docx acepta ambos en Document.save()."""
    doc = DocxDocument()

    seccion = doc.sections[0]
    seccion.left_margin = Cm(3)
    seccion.right_margin = Cm(2.5)
    seccion.top_margin = Cm(2.5)
    seccion.bottom_margin = Cm(2.5)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Times New Roman"
    estilo_normal.font.size = Pt(12)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.line_spacing = 1.5
    run_titulo = p_titulo.add_run(titulo.upper())
    run_titulo.bold = True
    run_titulo.font.name = "Times New Roman"
    run_titulo.font.size = Pt(14)

    for linea in cuerpo.split("\n"):
        linea = linea.strip()
        if not linea:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(linea)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    p_firma = doc.add_paragraph("_______________________________")
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.paragraph_format.line_spacing = 1.5

    p_nombre_firma = doc.add_paragraph("Nombre y firma")
    p_nombre_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_firma.paragraph_format.line_spacing = 1.5

    _agregar_numero_pagina(seccion)

    doc.save(str(destino) if isinstance(destino, Path) else destino)


def generar_nombre_archivo(tipo_nombre: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", tipo_nombre.lower()).strip("_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{slug}_{timestamp}"


# ---------------------------------------------------------------------------
# Historial
# ---------------------------------------------------------------------------

def guardar_historial(tipo_nombre: str, partes: list[str], archivo_base: str) -> None:
    historial = []
    if HISTORIAL_PATH.exists():
        try:
            historial = json.loads(HISTORIAL_PATH.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            historial = []

    historial.append(
        {
            "tipo": tipo_nombre,
            "fecha_creacion": datetime.now().isoformat(timespec="seconds"),
            "partes": partes,
            "archivo": archivo_base,
        }
    )

    try:
        HISTORIAL_PATH.write_text(
            json.dumps(historial, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except OSError as exc:
        print(f"No se pudo guardar el historial: {exc}")


# ---------------------------------------------------------------------------
# Flujo principal
# ---------------------------------------------------------------------------

def main() -> None:
    print("=" * 64)
    print("  REDACCION ASISTIDA DE DOCUMENTOS LEGALES")
    print("=" * 64)
    print(
        "Nota: los documentos generados son borradores de apoyo y no "
        "sustituyen la revision de un abogado antes de presentarlos ante "
        "una autoridad.\n"
    )

    try:
        provider_config = config.get_active_provider_config()
    except Exception as exc:
        print(f"No se pudo determinar el proveedor de LLM: {exc}")
        sys.exit(1)
    print(f"Proveedor de LLM activo: {provider_config['provider']} ({provider_config['model']})")

    tipo_key = elegir_tipo_documento()
    if tipo_key is None:
        print("Operacion cancelada.")
        return
    tipo_info = TIPOS_DOCUMENTO[tipo_key]

    print(f"\nVas a redactar: {tipo_info['nombre']}\n")
    datos = recolectar_datos(tipo_info)
    datos_formateados = formatear_datos_para_prompt(tipo_info, datos)

    prompt = construir_prompt(tipo_info["nombre"], tipo_info["enfoque"], datos_formateados)

    print("\nGenerando borrador con la IA...\n")
    borrador = generar_borrador(prompt, provider_config)
    if borrador is None:
        return

    print("-" * 64)
    print(borrador)
    print("-" * 64 + "\n")

    while True:
        respuesta = input("¿Quieres hacer cambios? (s/n): ").strip().lower()
        if respuesta == "s":
            instrucciones = pedir_texto_multilinea("Describe los cambios que quieres")
            prompt_cambios = construir_prompt_cambios(borrador, instrucciones)
            print("\nRegenerando documento con los cambios solicitados...\n")
            nuevo_borrador = generar_borrador(prompt_cambios, provider_config)
            if nuevo_borrador is None:
                continue
            borrador = nuevo_borrador
            print("-" * 64)
            print(borrador)
            print("-" * 64 + "\n")
        elif respuesta == "n":
            break
        else:
            print("Responde 's' o 'n'.")

    SALIDA_DIR.mkdir(parents=True, exist_ok=True)
    nombre_base = generar_nombre_archivo(tipo_info["nombre"])
    ruta_txt = SALIDA_DIR / f"{nombre_base}.txt"
    ruta_docx = SALIDA_DIR / f"{nombre_base}.docx"

    try:
        ruta_txt.write_text(borrador, encoding="utf-8")
        exportar_docx(ruta_docx, tipo_info["nombre"], borrador)
    except Exception as exc:
        print(f"Ocurrio un error al guardar los archivos: {exc}")
        return

    partes = extraer_partes(tipo_info, datos)
    guardar_historial(tipo_info["nombre"], partes, nombre_base)

    print(f"Documento guardado en:\n  {ruta_txt}\n  {ruta_docx}")
    print(f"Registro agregado a {HISTORIAL_PATH.name}")


if __name__ == "__main__":
    main()
